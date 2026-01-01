"""
文档处理器模块 - 负责读取和写入Word文档，保留格式和非文本元素
"""
from pathlib import Path
from typing import Callable, Optional
from dataclasses import dataclass
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from src.converter import ChineseConverter


@dataclass
class ProcessResult:
    """处理结果数据模型"""
    success: bool
    converted_chars: int
    error_message: str = ""


class DocumentProcessor:
    """Word文档处理器"""
    
    def __init__(self, converter: ChineseConverter):
        """
        初始化文档处理器
        
        Args:
            converter: 中文转换器实例
        """
        self._converter = converter
    
    def process_document(
        self, 
        input_path: Path, 
        output_path: Path,
        progress_callback: Optional[Callable[[int], None]] = None
    ) -> ProcessResult:
        """
        处理单个Word文档
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径
            progress_callback: 进度回调函数
            
        Returns:
            处理结果对象
        """
        try:
            # 确保是Path对象
            if isinstance(input_path, str):
                input_path = Path(input_path)
            if isinstance(output_path, str):
                output_path = Path(output_path)
            
            # 打开文档
            doc = Document(input_path)
            
            total_converted = 0
            
            # 计算总元素数用于进度
            total_elements = len(doc.paragraphs) + len(doc.tables)
            processed_elements = 0
            
            # 处理所有段落
            for paragraph in doc.paragraphs:
                converted = self._process_paragraph(paragraph)
                total_converted += converted
                processed_elements += 1
                if progress_callback and total_elements > 0:
                    progress_callback(int(processed_elements / total_elements * 100))
            
            # 处理所有表格
            for table in doc.tables:
                converted = self._process_table(table)
                total_converted += converted
                processed_elements += 1
                if progress_callback and total_elements > 0:
                    progress_callback(int(processed_elements / total_elements * 100))
            
            # 确保输出目录存在
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 保存文档
            doc.save(output_path)
            
            return ProcessResult(success=True, converted_chars=total_converted)
            
        except Exception as e:
            return ProcessResult(success=False, converted_chars=0, error_message=str(e))
    
    def _process_paragraph(self, paragraph: Paragraph) -> int:
        """
        处理单个段落，保留格式
        
        Args:
            paragraph: python-docx段落对象
            
        Returns:
            转换的字符数
        """
        converted_count = 0
        
        # 遍历段落中的每个run（文本块），保留格式
        for run in paragraph.runs:
            original_text = run.text
            if original_text:
                converted_text = self._converter.convert(original_text)
                # 计算转换的字符数
                for orig, conv in zip(original_text, converted_text):
                    if orig != conv:
                        converted_count += 1
                # 更新文本，保留run的格式（字体、大小、颜色等）
                run.text = converted_text
        
        return converted_count
    
    def _process_table(self, table: Table) -> int:
        """
        处理表格中的文本
        
        Args:
            table: python-docx表格对象
            
        Returns:
            转换的字符数
        """
        converted_count = 0
        
        # 遍历表格的每一行
        for row in table.rows:
            # 遍历每个单元格
            for cell in row.cells:
                # 处理单元格中的每个段落
                for paragraph in cell.paragraphs:
                    converted_count += self._process_paragraph(paragraph)
        
        return converted_count
