"""
DocumentProcessor单元测试
Feature: word-traditional-to-simplified
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.converter import ChineseConverter
from src.document_processor import DocumentProcessor, ProcessResult


@pytest.fixture
def converter():
    """创建转换器实例"""
    return ChineseConverter()


@pytest.fixture
def processor(converter):
    """创建文档处理器实例"""
    return DocumentProcessor(converter)


class TestDocumentProcessorUnit:
    """DocumentProcessor单元测试类"""
    
    def test_process_paragraph_traditional_text(self, processor, tmp_path):
        """测试段落文本转换 - Requirements 2.1"""
        # 创建测试文档
        doc = Document()
        doc.add_paragraph("這是繁體中文測試")
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        # 处理文档
        result = processor.process_document(input_path, output_path)
        
        # 验证结果
        assert result.success is True
        assert result.converted_chars > 0
        
        # 读取输出文档验证转换
        output_doc = Document(output_path)
        text = output_doc.paragraphs[0].text
        assert "这是" in text or "繁体" in text  # 验证转换发生
    
    def test_process_table_text(self, processor, tmp_path):
        """测试表格文本转换 - Requirements 2.2"""
        # 创建带表格的测试文档
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "繁體"
        table.cell(0, 1).text = "中文"
        table.cell(1, 0).text = "測試"
        table.cell(1, 1).text = "表格"
        
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        # 处理文档
        result = processor.process_document(input_path, output_path)
        
        # 验证结果
        assert result.success is True
        assert result.converted_chars > 0
        
        # 读取输出文档验证转换
        output_doc = Document(output_path)
        table = output_doc.tables[0]
        # 验证表格内容被转换
        assert table.cell(0, 0).text != "繁體" or table.cell(1, 0).text != "測試"
    
    def test_preserve_formatting(self, processor, tmp_path):
        """测试格式保留 - Requirements 2.5"""
        # 创建带格式的测试文档
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("繁體中文")
        run.bold = True
        run.italic = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        # 处理文档
        result = processor.process_document(input_path, output_path)
        
        # 验证结果
        assert result.success is True
        
        # 读取输出文档验证格式保留
        output_doc = Document(output_path)
        output_run = output_doc.paragraphs[0].runs[0]
        assert output_run.bold is True
        assert output_run.italic is True
        assert output_run.font.size == Pt(14)
        assert output_run.font.color.rgb == RGBColor(255, 0, 0)
    
    def test_preserve_paragraph_alignment(self, processor, tmp_path):
        """测试段落对齐格式保留"""
        doc = Document()
        para = doc.add_paragraph("繁體中文居中")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        result = processor.process_document(input_path, output_path)
        assert result.success is True
        
        output_doc = Document(output_path)
        assert output_doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    
    def test_process_empty_document(self, processor, tmp_path):
        """测试空文档处理"""
        doc = Document()
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        result = processor.process_document(input_path, output_path)
        
        assert result.success is True
        assert result.converted_chars == 0
    
    def test_process_simplified_only(self, processor, tmp_path):
        """测试纯简体文档（无需转换）"""
        doc = Document()
        doc.add_paragraph("这是简体中文")
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        result = processor.process_document(input_path, output_path)
        
        assert result.success is True
        assert result.converted_chars == 0
        
        output_doc = Document(output_path)
        assert output_doc.paragraphs[0].text == "这是简体中文"
    
    def test_process_mixed_content(self, processor, tmp_path):
        """测试混合内容（繁简混合）"""
        doc = Document()
        doc.add_paragraph("繁體和简体混合測試")
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        result = processor.process_document(input_path, output_path)
        
        assert result.success is True
        assert result.converted_chars > 0
    
    def test_process_nonexistent_file(self, processor, tmp_path):
        """测试不存在的文件"""
        input_path = tmp_path / "nonexistent.docx"
        output_path = tmp_path / "output.docx"
        
        result = processor.process_document(input_path, output_path)
        
        assert result.success is False
        assert result.error_message != ""
    
    def test_progress_callback(self, processor, tmp_path):
        """测试进度回调"""
        doc = Document()
        doc.add_paragraph("段落一")
        doc.add_paragraph("段落二")
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"
        doc.save(input_path)
        
        progress_values = []
        def callback(progress):
            progress_values.append(progress)
        
        result = processor.process_document(input_path, output_path, progress_callback=callback)
        
        assert result.success is True
        assert len(progress_values) > 0
        assert progress_values[-1] == 100  # 最后应该是100%
    
    def test_output_directory_creation(self, processor, tmp_path):
        """测试输出目录自动创建"""
        doc = Document()
        doc.add_paragraph("測試")
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "subdir" / "nested" / "output.docx"
        doc.save(input_path)
        
        result = processor.process_document(input_path, output_path)
        
        assert result.success is True
        assert output_path.exists()
