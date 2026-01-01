"""
集成测试 - 端到端文档转换流程和批量文件处理
Feature: word-traditional-to-simplified
Requirements: 2.1, 3.1
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.converter import ChineseConverter
from src.document_processor import DocumentProcessor
from src.conversion_manager import ConversionManager
from src.file_validator import FileValidator


class TestEndToEndConversion:
    """端到端文档转换流程测试 - Requirements 2.1"""
    
    @pytest.fixture
    def manager(self):
        """创建转换管理器实例"""
        return ConversionManager()
    
    def test_full_conversion_workflow(self, tmp_path):
        """测试完整的转换工作流程"""
        # 1. 创建测试文档
        input_path = tmp_path / "test_input.docx"
        doc = Document()
        doc.add_paragraph("這是一個繁體中文測試文檔")
        doc.add_paragraph("包含多個段落的內容")
        doc.save(input_path)
        
        # 2. 使用ConversionManager进行转换
        manager = ConversionManager()
        added = manager.add_files([input_path])
        assert len(added) == 1
        
        # 3. 执行转换
        tasks = manager.start_conversion()
        
        # 4. 验证结果
        assert len(tasks) == 1
        task = tasks[0]
        assert task.status == "completed"
        assert task.converted_chars > 0
        
        # 5. 验证输出文件
        output_path = manager.get_default_output_path(input_path)
        assert output_path.exists()
        
        # 6. 验证转换内容
        output_doc = Document(output_path)
        text = output_doc.paragraphs[0].text
        assert "这是" in text or "繁体" in text
    
    def test_conversion_preserves_formatting(self, tmp_path):
        """测试转换保留文档格式"""
        # 创建带格式的文档
        input_path = tmp_path / "formatted.docx"
        doc = Document()
        
        # 添加带格式的段落
        para = doc.add_paragraph()
        run = para.add_run("繁體中文")
        run.bold = True
        run.italic = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(255, 0, 0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(input_path)
        
        # 执行转换
        manager = ConversionManager()
        manager.add_files([input_path])
        tasks = manager.start_conversion()
        
        # 验证格式保留
        assert tasks[0].status == "completed"
        output_path = manager.get_default_output_path(input_path)
        output_doc = Document(output_path)
        
        output_para = output_doc.paragraphs[0]
        output_run = output_para.runs[0]
        
        assert output_run.bold is True
        assert output_run.italic is True
        assert output_run.font.size == Pt(16)
        assert output_run.font.color.rgb == RGBColor(255, 0, 0)
        assert output_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    
    def test_conversion_with_tables(self, tmp_path):
        """测试包含表格的文档转换"""
        input_path = tmp_path / "with_table.docx"
        doc = Document()
        
        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "繁體"
        table.cell(0, 1).text = "中文"
        table.cell(1, 0).text = "測試"
        table.cell(1, 1).text = "表格"
        
        doc.save(input_path)
        
        # 执行转换
        manager = ConversionManager()
        manager.add_files([input_path])
        tasks = manager.start_conversion()
        
        # 验证转换成功
        assert tasks[0].status == "completed"
        assert tasks[0].converted_chars > 0


class TestBatchFileProcessing:
    """批量文件处理测试 - Requirements 3.1"""
    
    def test_batch_conversion_multiple_files(self, tmp_path):
        """测试批量转换多个文件"""
        # 创建多个测试文档
        input_files = []
        for i in range(3):
            input_path = tmp_path / f"test_{i}.docx"
            doc = Document()
            doc.add_paragraph(f"這是第{i+1}個繁體文檔")
            doc.save(input_path)
            input_files.append(input_path)
        
        # 批量添加文件
        manager = ConversionManager()
        added = manager.add_files(input_files)
        assert len(added) == 3
        
        # 执行批量转换
        progress_calls = []
        def progress_callback(current, total, filename):
            progress_calls.append((current, total, filename))
        
        tasks = manager.start_conversion(progress_callback=progress_callback)
        
        # 验证所有文件都转换成功
        assert len(tasks) == 3
        for task in tasks:
            assert task.status == "completed"
            assert task.output_path.exists()
        
        # 验证进度回调被调用
        assert len(progress_calls) == 3
    
    def test_batch_conversion_with_output_dir(self, tmp_path):
        """测试批量转换到指定输出目录"""
        # 创建输入目录和输出目录
        input_dir = tmp_path / "input"
        output_dir = tmp_path / "output"
        input_dir.mkdir()
        output_dir.mkdir()
        
        # 创建测试文档
        input_files = []
        for i in range(2):
            input_path = input_dir / f"doc_{i}.docx"
            doc = Document()
            doc.add_paragraph(f"繁體文檔{i}")
            doc.save(input_path)
            input_files.append(input_path)
        
        # 执行转换到指定目录
        manager = ConversionManager()
        manager.add_files(input_files)
        tasks = manager.start_conversion(output_dir=output_dir)
        
        # 验证输出文件在指定目录
        for task in tasks:
            assert task.status == "completed"
            assert task.output_path.parent == output_dir
            assert task.output_path.exists()
    
    def test_batch_conversion_with_invalid_file(self, tmp_path):
        """测试批量转换包含无效文件的情况"""
        # 创建有效文档
        valid_path = tmp_path / "valid.docx"
        doc = Document()
        doc.add_paragraph("繁體中文")
        doc.save(valid_path)
        
        # 创建无效文件路径
        invalid_path = tmp_path / "nonexistent.docx"
        
        # 添加文件（包含无效文件）
        manager = ConversionManager()
        manager.add_files([valid_path, invalid_path])
        
        # 执行转换
        tasks = manager.start_conversion()
        
        # 验证有效文件成功，无效文件失败
        assert len(tasks) == 2
        
        valid_task = next(t for t in tasks if t.input_path == valid_path)
        invalid_task = next(t for t in tasks if t.input_path == invalid_path)
        
        assert valid_task.status == "completed"
        assert invalid_task.status == "failed"
        assert "不存在" in invalid_task.error_message
    
    def test_file_deduplication_in_batch(self, tmp_path):
        """测试批量添加时的文件去重"""
        # 创建测试文档
        input_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("測試")
        doc.save(input_path)
        
        # 多次添加同一文件
        manager = ConversionManager()
        added1 = manager.add_files([input_path])
        added2 = manager.add_files([input_path])
        added3 = manager.add_files([input_path])
        
        # 验证去重
        assert len(added1) == 1
        assert len(added2) == 0
        assert len(added3) == 0
        assert len(manager.get_files()) == 1
    
    def test_clear_and_reconvert(self, tmp_path):
        """测试清空后重新转换"""
        # 创建测试文档
        input_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("繁體中文")
        doc.save(input_path)
        
        manager = ConversionManager()
        
        # 第一次转换
        manager.add_files([input_path])
        tasks1 = manager.start_conversion()
        assert tasks1[0].status == "completed"
        
        # 清空并重新添加
        manager.clear_files()
        assert len(manager.get_files()) == 0
        
        # 第二次转换
        manager.add_files([input_path])
        tasks2 = manager.start_conversion()
        assert tasks2[0].status == "completed"


class TestComponentIntegration:
    """组件集成测试"""
    
    def test_converter_processor_integration(self, tmp_path):
        """测试Converter和Processor的集成"""
        # 创建测试文档
        input_path = tmp_path / "test.docx"
        output_path = tmp_path / "output.docx"
        
        doc = Document()
        doc.add_paragraph("繁體中文測試")
        doc.save(input_path)
        
        # 直接使用Converter和Processor
        converter = ChineseConverter()
        processor = DocumentProcessor(converter)
        
        result = processor.process_document(input_path, output_path)
        
        assert result.success is True
        assert result.converted_chars > 0
        assert output_path.exists()
    
    def test_validator_manager_integration(self, tmp_path):
        """测试Validator和Manager的集成"""
        # 创建有效文档
        valid_path = tmp_path / "valid.docx"
        doc = Document()
        doc.add_paragraph("測試")
        doc.save(valid_path)
        
        # 创建无效文件
        invalid_path = tmp_path / "invalid.txt"
        invalid_path.write_text("not a word doc")
        
        # 验证文件
        valid_result = FileValidator.validate(valid_path)
        invalid_result = FileValidator.validate(invalid_path)
        
        assert valid_result[0] is True
        assert invalid_result[0] is False
        
        # Manager应该能处理有效文件
        manager = ConversionManager()
        manager.add_files([valid_path])
        tasks = manager.start_conversion()
        
        assert tasks[0].status == "completed"
